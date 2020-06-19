import database as db
import urllib.request
from datetime import date
from datetime import timedelta
import pandas as pd
import os
import functools
import requests
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Inches

#was ist das ziel_-° am ende struktur mit csv datei aus der author-dict eingelesen
#brauche write author-dict, read author-dict, initialize-function mit initialem schreiben und funktion die das returnt, sowie funktion, die anhang von liste fon seiten das erweitert (get_full_author_dict: dort nicht initialisierung)

#class/file for sitemap imports
#wenn import: wie verhindern dass alle alten werte nciht stärker 
#last data definieren, wenn gleich: dann überschreiben, sonst ergänzen

#Generiert das vollen author_dict ausgehend von einer Liste an gescrapten Seiten (eingetragen in wiki_index.txt)
"""
def get_full_author_dict_from_wiki(index_file):
	list_of_pages = get_scraped_wiki_pages(index_file)
	list_of_names = list(map(get_name, list_of_pages)) #schöne Namen ausgehend von den Links
	author_dict = dict()
	for i in range(0, len(list_of_pages)):#Geht durch jeden Artikel durch und ergänzt das author_dict
		author_dict = get_author_frame(list_of_names[i], list_of_pages[i], author_dict)
	return author_dict
"""

#Generiert ein author_dict neu anhand eines gescrapten Artikels, oder updated ein bestehendes
#=dict() rausnemhen, kriegt dataframe direct übergebe
def get_author_frame_article(revision_json, author_frame):
	for edit in revision_json:
		actualize_cell_value(edit["user"], pd.to_datetime(edit["timestamp"]).normalize().tz_localize(tz=None), author_frame)
				

#author_name as string, date as pandas datetime
def actualize_cell_value(author_name, date, author_frame):
	if author_name not in author_frame:
		author_frame[author_name] = 0
		author_frame.loc[date, author_name] = 1
	else: 
		if date in author_frame.index:
			if pd.notna(author_frame.loc[date, author_name]):
				author_frame.loc[date, author_name] = author_frame.loc[date, author_name] + 1
			else:
				author_frame.loc[date, author_name] = 1
		else:
			author_frame.loc[date, author_name] = 1
	return author_frame

#limit_date as python datetime
def init_author_frame(topic, request_session=None):
	today = pd.to_datetime(date.today())
	author_frame = pd.DataFrame([today], columns=['date'])
	author_frame['date'] = pd.to_datetime(author_frame['date'])
	author_frame = author_frame.set_index('date')

	if request_session is None:
		request_session = requests.Session()

	sitemap_pages = db.scrape_sitemap(request_session, topic)[topic]
	print("Generating author-frame for topic: %s"%topic)

	for page in sitemap_pages:
		get_author_frame_article(db.get_article_revisions(page, request_session), author_frame)

	return author_frame.fillna(0)


def create_from_zero(*topics, request_session=None):
	if request_session is None:
		request_session = requests.Session()

	topic_frames = {topic: init_author_frame(topic, request_session=request_session) for topic in topics}

	mfnf_frame = functools.reduce(lambda a,b: a.add(b, fill_value=0), topic_frames.values())

	os.makedirs("topic_frames", exist_ok=True)

	for topic, frame in topic_frames.items():
		write_author_frame(frame, "topic_frames/"+topic.replace(" ", "_")+".csv")

	write_author_frame(mfnf_frame, "topic_frames/author_frame.csv")

	return topic_frames, mfnf_frame

def create_mfnf_frame(topic_frames):
	return functools.reduce(lambda a,b: a.add(b, fill_value=0), topic_frames.values())

#Annahme topic_frames auf selben zeitlichen Stand wie mfnf frame
def actualize_author_frames(topic_frames, mfnf_frame=None, request_session=None):
	if request_session is None:
		request_session = requests.Session()

	sitemap_pages = db.scrape_sitemap(request_session, *(topic_frames.keys()))

	if mfnf_frame is not None:
		mfnf_frame.sort_index(inplace=True)
		mfnf_frame.drop(mfnf_frame.index.max(), inplace=True)

	for topic in topic_frames.keys():
		print("Actualizing topic %s"%topic)

		topic_frames[topic].sort_index(inplace=True)
		last_date = topic_frames[topic].index.max()
		topic_frames[topic].drop(last_date, inplace=True)

		for page in sitemap_pages[topic]:
			revisions = db.get_article_revisions(page, request_session, start_date=last_date)
			if revisions:
				get_author_frame_article(revisions, topic_frames[topic]) 
				if mfnf_frame is not None:
					get_author_frame_article(revisions, mfnf_frame)


def write_author_frame(author_frame, filename):
	author_frame.to_csv(filename, mode="w+")

def import_author_frame(filename="topic_frames/author_frame.csv"):
	df = pd.read_csv(filename, na_values=[" "], parse_dates=["date"], index_col="date")
	return df

def read(*topics, mfnf_frame=True):
	existing=[]
	topic_frames = {}
	for topic in topics:
		if not os.path.isfile("topic_frames/"+topic.replace(" ", "_")+".csv"):
			topic_frames[topic] = init_author_frame(topic)
			write_author_frame(topic_frames[topic], "topic_frames/"+topic.replace(" ", "_")+".csv")
		else:
			existing.append(topic)
	if existing:
		for topic in existing:
			topic_frames[topic] = import_author_frame("topic_frames/"+topic.replace(" ", "_")+".csv")

	if mfnf_frame:
		mfnf_frame = import_author_frame("topic_frames/author_frame.csv")
		return topic_frames, mfnf_frame
	else:
		return topic_frames

def write(topic_frames, mfnf=None):
	os.makedirs("topic_frames", exist_ok=True)
	for topic in topic_frames.keys():	
		write_author_frame(topic_frames[topic], "topic_frames/"+topic.replace(" ", "_")+".csv")
	if mfnf is not None:
		write_author_frame(mfnf, "topic_frames/author_frame.csv")

def create_accumulation(topic_frame, days=90):
	#topic_frame.sort_index(inplace=True)
	topic_frame.index = pd.DatetimeIndex(topic_frame.index)
	index = pd.date_range(topic_frame.index.min(), pd.to_datetime(date.today()))
	topic_frame = topic_frame.reindex(index, fill_value=0)
	return topic_frame.rolling(min_periods=1, window=days).sum()


def get_authors(accumulated_frame, edits, datum=date.today()):
	authors = list(accumulated_frame.columns[accumulated_frame.loc[pd.to_datetime(datum)]>edits])
	edits = list(accumulated_frame[authors].loc[pd.to_datetime(datum)])
	author_dict =  dict(sorted(list(zip(authors,edits)), key=lambda x: x[1], reverse=True))  
	if "authors" in author_dict.keys():
		del author_dict["authors"]
	if "active authors" in author_dict.keys():
		del author_dict["active authors"]
	return author_dict

def report(topics, request_session=None, initialize=False, actualize=True):
	if request_session is None:
		request_session = requests.Session()

	if initialize:
		create_from_zero(*topics)

	topic_frames, mfnf = read(*topics)
	
	if actualize:
		actualize_author_frames(topic_frames, mfnf)
		write(topic_frames, mfnf)
	
	print("Accumulating edits")
	accumulation  = create_accumulation(mfnf)
	limit_date = pd.to_datetime(date.today()-timedelta(days=365))

	#Liste aktiver Autor*innen
	print("Generating list of active authors")
	active_authors = get_authors(accumulation, 9)

	accumulation["authors"] = accumulation.apply(lambda data_row: len(data_row[data_row>0]), axis=1)
	accumulation["active authors"] = accumulation.apply(lambda data_row: len(data_row[data_row>9]), axis=1)

	os.makedirs("plots", exist_ok=True)

	#Historische Gesamtschauen
	print("Generating historical plots for authors and active authors")
	plt.plot(accumulation.index, accumulation["authors"], label="Autor*innen")
	plt.plot(accumulation.index, accumulation["active authors"], label="Aktive Autor*innen")
	plt.legend()
	plt.savefig("plots/all_years.png")
	plt.close()

	accumulation = accumulation[accumulation.index > limit_date]
	accumulation["percent"] = accumulation["active authors"]/accumulation["authors"]

	#Aktuellere Plots
	print("Generating last years plots for authors and active authors")
	plt.plot(accumulation.index, accumulation["authors"], label="Autor*innen")
	plt.plot(accumulation.index, accumulation["active authors"], label="Aktive Autor*innen")
	plt.legend()
	plt.savefig("plots/last_year.png")
	plt.close()

	#Verhältnis aktive Autor*innen zu anderen:
	print("Generating plot for percentage of active editors")
	plt.plot(accumulation.index, accumulation["percent"], label="Prozent")
	plt.legend()
	plt.savefig("plots/percent.png")
	plt.close()


	#Topic Plots
	print("Generating plots for development of authors and active authors in topics")
	topic_accumulations = {topic: create_accumulation(topic_frames[topic]) for topic in topic_frames.keys()}
	for topic in topic_accumulations:
		topic_accumulations[topic]["authors"] = topic_accumulations[topic].apply(lambda data_row: len(data_row[data_row>0]), axis=1)
		topic_accumulations[topic]["active authors"] = topic_accumulations[topic].apply(lambda data_row: len(data_row[data_row>9]), axis=1)

		plt.plot(topic_accumulations[topic][topic_accumulations[topic].index>limit_date].index, topic_accumulations[topic][topic_accumulations[topic].index>limit_date]["authors"], label="Autor*innen %s"%topic)
		plt.plot(topic_accumulations[topic][topic_accumulations[topic].index>limit_date].index, topic_accumulations[topic][topic_accumulations[topic].index>limit_date]["active authors"], label="Aktive Autor*innen %s"%topic)
	plt.savefig("plots/topics_last_year.png")
	plt.legend()
	plt.close()

	#Edit Plot
	print("Generating historical overview of edits")
	mfnf = mfnf[mfnf.index>limit_date]
	mfnf["edits"] = mfnf.sum(axis=1)

	#Filling empty days
	mfnf.index = pd.DatetimeIndex(mfnf.index)
	index = pd.date_range(mfnf.index.min(), pd.to_datetime(date.today()))
	mfnf = mfnf.reindex(index, fill_value=0)

	plt.plot(mfnf.index, mfnf["edits"], label="Edits")
	plt.savefig("plots/edits.png")
	plt.legend()
	plt.close()

	print("Generating report")
	document = Document()

	document.add_heading("Report MfNF", 0)
	document.add_paragraph("Aktive Autor*innen: %s"%len(active_authors.keys()))

	for author in active_authors.keys():
		document.add_paragraph("%s: %s\n"%(author, active_authors[author]), style='List Number')

	document.add_paragraph("\n")
	document.add_heading('Autor*innenentwicklung über die Jahre', level=3)	
	document.add_picture('plots/all_years.png', width=Inches(4))

	document.add_heading('Autor*innenentwicklung letztes Jahr (bis heute)', level=3)	
	document.add_picture('plots/last_year.png', width=Inches(4))

	document.add_heading('Prozent aktiver Autor*innen letztes Jahr (bis heute)', level=3)	
	document.add_picture('plots/percent.png', width=Inches(4))

	document.add_heading('Autor*innenentwicklung Fächer', level=3)	
	document.add_picture('plots/topics_last_year.png', width=Inches(4))

	document.add_heading("Autor*innen mit über 5 edits Fächer:", level=3)

	for topic in topic_accumulations.keys():
		topic_accumulations[topic]
		authors = get_authors(topic_accumulations[topic], 4)
		document.add_paragraph("%s: %s"%(topic, len(authors.keys())))
		for author in authors:
			document.add_paragraph("%s: %s"%(author, authors[author]), style='List Bullet')

	document.add_page_break()

	document.save('report.docx')


if __name__ == "__main__":
	topics = ["Grundlagen der Mathematik", "Analysis 1", "Lineare Algebra 1","Maßtheorie","Real Analysis", "Mitmachen für (Nicht-)Freaks"] #, ["Buchanfänge", "sitemap_files/buch_sitemap.html", "index_files/buch_index.txt"], ["Mitmachen für (Nicht-)Freaks", "sitemap_files/mitm_sitemap.html", "index_files/mitm_index.txt"]
	S = requests.Session()

	create_from_zero(*topics)
	topic_frames, mfnf = read(*topics)
	actualize_author_frames(topic_frames, mfnf)
	write(topic_frames, mfnf)

	limit_date = pd.to_datetime(date.today()-timedelta(days=365))
	mfnf = mfnf[mfnf.index>limit_date]
	accumulation  = create_accumulation(mfnf)




#accumulation["authors"] = df[df.columns[df.loc["2020-06-15"]>9]].loc["2020-06-15"]

#print(db.get_section_id(S, "Analysis 1"))

#print(db.scrape_sitemap_topic(S, 2))
#df = init_author_frame("Analysis 1")
#topics, mfnf_frame = create_from_zero(*topics)
#neue Datenquelle und standardisierung als json-daten wie ausgegeben --> dann ist das bauen dieser boards nur eine funktionalität
"""
PARAMS = {
    "action": "query",
    "prop": "revisions",
    "titles": "Mathe_für_Nicht-Freaks:_Sitemap",
    "rvlimit": "max",
    "rvprop": "timestamp|user",
    "rvdir": "newer",
    "rvstart": "2018-07-01T00:00:00Z",
    "rvslots": "main",
    "formatversion": "2",
    "format": "json"
}
"""
#todo: überlegen was ich visualierien kann und ausgeben --> kathi fragen
#legenden
#todo: cronjob und wann ausführen/wann aktualisieren

#todo: requests von aussen