import urllib.request
from datetime import date
from datetime import timedelta
import pandas as pd
import os
import requests
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Inches
import authors_MfNF as am


def report(topics, request_session=None, initialize=False, actualize=True):
	if request_session is None:
		request_session = requests.Session()

	if initialize:
		am.create_from_zero(*topics)

	topic_frames, mfnf = am.read(*topics)
	
	if actualize:
		am.actualize_author_frames(topic_frames, mfnf)
		am.write(topic_frames, mfnf)
	
	print("Accumulating edits")
	accumulation  = am.create_accumulation(mfnf)
	limit_date = pd.to_datetime(date.today()-timedelta(days=365))

	#Liste aktiver Autor*innen
	print("Generating list of active authors")
	active_authors =  am.get_authors(accumulation, 9)

	accumulation["authors"] = accumulation.apply(lambda data_row: len(data_row[data_row>0]), axis=1)
	accumulation["active authors"] = accumulation.apply(lambda data_row: len(data_row[data_row>9]), axis=1)

	os.makedirs("plots", exist_ok=True)

	#Historische Gesamtschauen
	print("Generating historical plots for authors and active authors")
	plt.plot(accumulation.index, accumulation["authors"], label="Autor*innen")
	plt.plot(accumulation.index, accumulation["active authors"], label="Aktive Autor*innen")
	plt.legend()
	plt.savefig("plots/all_years.png")
	plt.close()

	accumulation = accumulation[accumulation.index > limit_date]
	accumulation["percent"] = accumulation["active authors"]/accumulation["authors"]

	#Aktuellere Plots
	print("Generating last years plots for authors and active authors")
	plt.plot(accumulation.index, accumulation["authors"], label="Autor*innen")
	plt.plot(accumulation.index, accumulation["active authors"], label="Aktive Autor*innen")
	plt.legend()
	plt.savefig("plots/last_year.png")
	plt.close()

	#Verhältnis aktive Autor*innen zu anderen:
	print("Generating plot for percentage of active editors")
	plt.plot(accumulation.index, accumulation["percent"], label="Prozent aktiver Autor*innen")
	plt.legend()
	plt.savefig("plots/percent.png")
	plt.close()


	#Topic Plots
	print("Generating plots for development of authors and active authors in topics")
	topic_accumulations = {topic:  am.create_accumulation(topic_frames[topic]) for topic in topic_frames.keys()}
	for topic in topic_accumulations:
		topic_accumulations[topic]["authors"] = topic_accumulations[topic].apply(lambda data_row: len(data_row[data_row>0]), axis=1)
		topic_accumulations[topic]["active authors"] = topic_accumulations[topic].apply(lambda data_row: len(data_row[data_row>9]), axis=1)

		plt.plot(topic_accumulations[topic][topic_accumulations[topic].index>limit_date].index, topic_accumulations[topic][topic_accumulations[topic].index>limit_date]["authors"], label="Autor*innen %s"%topic)
		plt.plot(topic_accumulations[topic][topic_accumulations[topic].index>limit_date].index, topic_accumulations[topic][topic_accumulations[topic].index>limit_date]["active authors"], label="Aktive Autor*innen %s"%topic)
	plt.legend(fontsize="xx-small")
	plt.savefig("plots/topics_last_year.png")
	plt.close()

	#Edit Plot
	print("Generating historical overview of edits")
	mfnf = mfnf[mfnf.index>limit_date]
	mfnf["edits"] = mfnf.sum(axis=1)

	#Filling empty days
	mfnf.index = pd.DatetimeIndex(mfnf.index)
	index = pd.date_range(mfnf.index.min(), pd.to_datetime(date.today()))
	mfnf = mfnf.reindex(index, fill_value=0)

	plt.plot(mfnf.index, mfnf["edits"], label="Edits")
	plt.savefig("plots/edits.png")
	plt.legend()
	plt.close()

	print("Generating report")
	document = Document()

	document.add_heading("Report MfNF", 0)
	document.add_paragraph("Aktive Autor*innen: %s"%len(active_authors.keys()))

	for author in active_authors.keys():
		document.add_paragraph("%s: %s\n"%(author, active_authors[author]), style='List Number')

	document.add_paragraph("\n")
	document.add_heading('Autor*innenentwicklung über die Jahre', level=3)	
	document.add_picture('plots/all_years.png', width=Inches(4))

	document.add_heading('Autor*innenentwicklung letztes Jahr (bis heute)', level=3)	
	document.add_picture('plots/last_year.png', width=Inches(4))

	document.add_heading('Prozent aktiver Autor*innen letztes Jahr (bis heute)', level=3)	
	document.add_picture('plots/percent.png', width=Inches(4))

	document.add_heading('Autor*innenentwicklung Fächer', level=3)	
	document.add_picture('plots/topics_last_year.png', width=Inches(4))

	document.add_heading("Autor*innen mit über 5 edits Fächer:", level=3)

	for topic in topic_accumulations.keys():
		authors =  am.get_authors(topic_accumulations[topic], 4)
		document.add_paragraph("%s: %s"%(topic, len(authors.keys())))
		for author in authors:
			document.add_paragraph("%s: %s"%(author, authors[author]), style='List Bullet')

	document.add_page_break()
	os.remove("report.docx")
	document.save('report.docx')


if __name__ == "__main__":
	topics = ["Grundlagen der Mathematik", "Analysis 1", "Lineare Algebra 1","Maßtheorie","Real Analysis", "Mitmachen für (Nicht-)Freaks"] #, ["Buchanfänge", "sitemap_files/buch_sitemap.html", "index_files/buch_index.txt"], ["Mitmachen für (Nicht-)Freaks", "sitemap_files/mitm_sitemap.html", "index_files/mitm_index.txt"]
	"""S = requests.Session()

	create_from_zero(*topics)
	topic_frames, mfnf = read(*topics)
	actualize_author_frames(topic_frames, mfnf)
	write(topic_frames, mfnf)

	limit_date = pd.to_datetime(date.today()-timedelta(days=365))
	mfnf = mfnf[mfnf.index>limit_date]
	accumulation  = create_accumulation(mfnf)"""

	if not os.path.isdir("topic_frames"):
		report(topics, initialize=True, actualize=False)
	else:
		report(topics)
